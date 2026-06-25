import os
import docx
from src.tools.file_readers import read_resume_docx, read_github_json
from src.graph import build_resume_to_portfolio_graph
from loguru import logger
from dotenv import load_dotenv


def save_insights_to_docx(insights: str, output_path: str) -> bool:
    """
        Saves the generated LLM insights into a styled Word Document.
    """
    try:
        # Ensure the output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok = True)
        
        doc = docx.Document()
        doc.add_heading('🚀 Resume to Portfolio Insights 🚀', 0)
        
        # Split insights by double newlines to create separate paragraphs
        for paragraph in insights.split('\n\n'):
            doc.add_paragraph(paragraph.strip())
            
        doc.save(output_path)
        logger.success(f"Insights successfully written to document: {output_path}")

        return True
    except Exception as e:
        logger.error(f"Failed to save insights document. Error: {e}")
        return False

def main():
    # Ensure environment variables (API keys) are loaded
    load_dotenv()

    # Define local file paths
    # Ensure your unstructured docx and json files are in these exact locations
    resume_path = "data/raw/John Doe - Resume.docx"
    github_path = "data/raw/GitHub REST API (Deep JSON Payload).json"
    output_docx_path = "data/output/Career_Coach_Insights.docx"

    logger.info("Starting the Career Coach Agent application...")

    # Step 1 & 2: Fetch the local unstructured and structured data
    resume_text = read_resume_docx(resume_path)
    github_data = read_github_json(github_path)

    # Guardrails to ensure we have data before invoking the LLM
    if not resume_text:
        logger.error("Resume text is empty. Please check the file path and contents.")
        return
    
    if not github_data or "error" in github_data:
        logger.error("GitHub data is missing or invalid. Please check the JSON payload.")
        return

    # Initialize the starting state for LangGraph
    initial_state = {
        "resume_text": resume_text,
        "github_data": github_data,
        "extracted_resume_skills": [],
        "insights": ""
    }

    # Step 3: Compile and run the logic engine
    graph = build_resume_to_portfolio_graph()
    
    logger.info("Invoking the LangGraph workflow. Please wait...")
    
    # Run the graph and capture the final state and save it in a document
    final_state = graph.invoke(initial_state)
    save_insights_to_docx(final_state["insights"], output_docx_path)
    logger.success("Execution completed successfully. You can view your insights at file location: {output_docx_path}")

if __name__ == "__main__":
    main()